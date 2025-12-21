from __future__ import annotations

import base64
import hashlib
import json
import logging
import os
import re
import time
import urllib.parse
import urllib.request
import urllib.error
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field

from app.core.config import settings
from app.db.models import ParentChunk
from app.db.repository import Repository
from app.schemas.ingest import IngestRequest, IngestResponse
from app.services.chunking.child_chunker import ChildChunkConfig, make_children
from app.services.embeddings.e5 import E5Embedder
from app.services.vectorstore.qdrant import QdrantStore

log = logging.getLogger(__name__)
router = APIRouter()


# --------------------------------------------------------------------------------------
# Basic helpers
# --------------------------------------------------------------------------------------

def _truthy(v: str | None) -> bool:
    if v is None:
        return False
    return v.strip().lower() in {"1", "true", "yes", "y", "on"}


def _stable_id(prefix: str, b: bytes) -> str:
    h = hashlib.sha1(b).hexdigest()
    return f"{prefix}_{h}"


def _guess_doc_meta(text: str, fallback_title: str) -> Tuple[str, str, str]:
    title = fallback_title
    doc_number = ""
    doc_date = ""

    m = re.search(r"(№|N)\s*([0-9A-Za-z\-/]+)", text)
    if m:
        doc_number = m.group(2)

    d = re.search(r"(\d{2}\.\d{2}\.\d{4})", text)
    if d:
        doc_date = d.group(1)

    for line in text.splitlines():
        line = line.strip()
        if line:
            title = line[:160]
            break

    return title, doc_number, doc_date


def _extract_text_from_docx(fp: Path) -> str:
    from docx import Document  # type: ignore

    doc = Document(str(fp))
    parts: List[str] = []
    last = ""

    def push(s: str) -> None:
        nonlocal last
        s = (s or "").strip()
        if not s or s == last:
            return
        parts.append(s)
        last = s

    for p in doc.paragraphs:
        push(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                push(cell.text)

    return "\n".join(parts).strip()


def _parent_chunk(text: str, parent_chars: int) -> List[Tuple[int, str]]:
    if not text:
        return []
    out: List[Tuple[int, str]] = []
    part_no = 1
    start = 0
    n = len(text)
    while start < n:
        chunk = text[start: start + parent_chars].strip()
        if chunk:
            out.append((part_no, chunk))
            part_no += 1
        start += parent_chars
    return out


def _build_children(parents: List[ParentChunk], child_cfg: ChildChunkConfig) -> List[Any]:
    workers = int(os.getenv("CHILD_WORKERS", "1"))
    workers = max(1, min(workers, 32))
    log.warning("[ingest] child_workers=%d", workers)

    if workers == 1:
        children: List[Any] = []
        for i, p in enumerate(parents, start=1):
            children.extend(make_children(p, child_cfg))
            if i % 25 == 0:
                log.warning(
                    "[ingest] children_build_progress parents=%d/%d children=%d",
                    i, len(parents), len(children),
                )
        return children

    children: List[Any] = []
    with ThreadPoolExecutor(max_workers=workers) as ex:
        futures = [ex.submit(make_children, p, child_cfg) for p in parents]
        done = 0
        for f in as_completed(futures):
            children.extend(f.result())
            done += 1
            if done % 25 == 0:
                log.warning(
                    "[ingest] children_build_progress parents=%d/%d children=%d",
                    done, len(parents), len(children),
                )
    return children


# --------------------------------------------------------------------------------------
# Qdrant singleton
# --------------------------------------------------------------------------------------

_qdrant_store: QdrantStore | None = None


def get_qdrant_store() -> QdrantStore:
    global _qdrant_store
    if _qdrant_store is None:
        embedder = E5Embedder()
        _qdrant_store = QdrantStore(embedder=embedder)
    return _qdrant_store


# --------------------------------------------------------------------------------------
# Elastic via REST (so we can print bulk error samples)
# --------------------------------------------------------------------------------------

@dataclass(frozen=True)
class ElasticConfig:
    url: str
    index: str
    user: str | None
    password: str | None
    timeout_s: float
    refresh: bool
    recreate: bool
    bulk_actions: int
    bulk_error_samples: int


def _elastic_cfg() -> ElasticConfig:
    url = os.getenv("ELASTIC_URL") or getattr(settings, "ELASTIC_URL", "") or "http://localhost:9200"
    index = os.getenv("ELASTIC_INDEX") or getattr(settings, "ELASTIC_INDEX", "") or "parent_chunks"

    return ElasticConfig(
        url=url.rstrip("/"),
        index=index,
        user=os.getenv("ELASTIC_USER") or None,
        password=os.getenv("ELASTIC_PASS") or None,
        timeout_s=float(os.getenv("ELASTIC_TIMEOUT_S", "30")),
        refresh=_truthy(os.getenv("ELASTIC_REFRESH", "0")),
        recreate=_truthy(os.getenv("ELASTIC_RECREATE", "0")),
        bulk_actions=max(50, min(int(os.getenv("ELASTIC_BULK_ACTIONS", "500")), 5000)),
        bulk_error_samples=max(1, min(int(os.getenv("ELASTIC_BULK_ERROR_SAMPLES", "3")), 50)),
    )


def _elastic_headers(cfg: ElasticConfig, ndjson: bool = False) -> Dict[str, str]:
    headers = {"Accept": "application/json"}
    headers["Content-Type"] = "application/x-ndjson" if ndjson else "application/json; charset=utf-8"
    if cfg.user and cfg.password:
        token = base64.b64encode(f"{cfg.user}:{cfg.password}".encode("utf-8")).decode("ascii")
        headers["Authorization"] = f"Basic {token}"
    return headers


def _http_json(method: str, url: str, *, headers: Dict[str, str], body: Any | None, timeout_s: float) -> Tuple[int, Any]:
    data: bytes | None = None
    if body is not None:
        data = json.dumps(body, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(url=url, data=data, method=method, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=timeout_s) as r:
            raw = r.read()
            ct = (r.headers.get("Content-Type") or "").lower()
            if raw and "application/json" in ct:
                return r.status, json.loads(raw.decode("utf-8"))
            if raw:
                try:
                    return r.status, json.loads(raw.decode("utf-8"))
                except Exception:
                    return r.status, raw.decode("utf-8", errors="replace")
            return r.status, None
    except urllib.error.HTTPError as e:
        raw = e.read()
        try:
            parsed = json.loads(raw.decode("utf-8")) if raw else None
        except Exception:
            parsed = raw.decode("utf-8", errors="replace") if raw else None
        return e.code, parsed
    except Exception as e:
        raise RuntimeError(f"Elastic request failed: {method} {url}: {e}") from e


def _http_ndjson(method: str, url: str, *, headers: Dict[str, str], ndjson_body: bytes, timeout_s: float) -> Tuple[int, Any]:
    req = urllib.request.Request(url=url, data=ndjson_body, method=method, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=timeout_s) as r:
            raw = r.read()
            return r.status, json.loads(raw.decode("utf-8")) if raw else None
    except urllib.error.HTTPError as e:
        raw = e.read()
        try:
            parsed = json.loads(raw.decode("utf-8")) if raw else None
        except Exception:
            parsed = raw.decode("utf-8", errors="replace") if raw else None
        return e.code, parsed
    except Exception as e:
        raise RuntimeError(f"Elastic request failed: {method} {url}: {e}") from e


def _elastic_ensure_index(cfg: ElasticConfig) -> None:
    index_url = f"{cfg.url}/{urllib.parse.quote(cfg.index)}"

    if cfg.recreate:
        st, _ = _http_json("DELETE", index_url, headers=_elastic_headers(cfg), body=None, timeout_s=cfg.timeout_s)
        if st not in (200, 202, 404):
            log.warning("[ingest] elastic_delete_index_unexpected status=%s", st)

    st, _ = _http_json("GET", index_url, headers=_elastic_headers(cfg), body=None, timeout_s=cfg.timeout_s)
    if st == 200:
        return
    if st != 404:
        raise RuntimeError(f"Elastic index check failed: status={st}")

    mapping = {
        "mappings": {
            "properties": {
                "parent_id": {"type": "keyword"},
                "doc_id": {"type": "keyword"},
                "part_no": {"type": "integer"},
                "title": {"type": "text", "fields": {"keyword": {"type": "keyword", "ignore_above": 256}}},
                "section_path": {"type": "keyword"},
                "anchor": {"type": "keyword"},
                "content": {"type": "text"},
            }
        }
    }
    st, resp = _http_json("PUT", index_url, headers=_elastic_headers(cfg), body=mapping, timeout_s=cfg.timeout_s)
    if st not in (200, 201):
        raise RuntimeError(f"Elastic create index failed: status={st} resp={resp}")


def _elastic_refresh(cfg: ElasticConfig) -> None:
    st, resp = _http_json(
        "POST",
        f"{cfg.url}/{urllib.parse.quote(cfg.index)}/_refresh",
        headers=_elastic_headers(cfg),
        body=None,
        timeout_s=cfg.timeout_s,
    )
    if st not in (200, 201):
        log.warning("[ingest] elastic_refresh_failed status=%s resp=%s", st, resp)


def _elastic_delete_by_doc_id(cfg: ElasticConfig, doc_id: str) -> int:
    url = f"{cfg.url}/{urllib.parse.quote(cfg.index)}/_delete_by_query"
    body = {
        "query": {
            "bool": {
                "should": [
                    {"term": {"doc_id": doc_id}},
                    {"term": {"doc_id.keyword": doc_id}},
                ],
                "minimum_should_match": 1,
            }
        }
    }
    st, resp = _http_json("POST", url, headers=_elastic_headers(cfg), body=body, timeout_s=cfg.timeout_s)
    if st not in (200, 201):
        log.warning("[ingest] elastic_delete_by_query_failed doc_id=%s status=%s resp=%s", doc_id, st, resp)
        return 0
    try:
        return int((resp or {}).get("deleted", 0))
    except Exception:
        return 0


def _elastic_bulk_index_parents(
    cfg: ElasticConfig,
    parents: List[ParentChunk],
    *,
    refresh_wait: bool,
) -> Tuple[int, int, List[Dict[str, Any]]]:
    if not parents:
        return 0, 0, []

    bulk_url = f"{cfg.url}/_bulk"
    if refresh_wait:
        bulk_url += "?refresh=wait_for"  # мгновенная видимость без ручного ожидания

    ok_total = 0
    err_total = 0
    samples: List[Dict[str, Any]] = []

    step = cfg.bulk_actions
    for start in range(0, len(parents), step):
        batch = parents[start: start + step]
        lines: List[str] = []
        for p in batch:
            meta = {"index": {"_index": cfg.index, "_id": p.parent_id}}
            doc = {
                "parent_id": p.parent_id,
                "doc_id": p.doc_id,
                "part_no": int(p.part_no),
                "title": p.title or "",
                "section_path": p.section_path or "",
                "anchor": p.anchor or "",
                "content": p.content or "",
            }
            lines.append(json.dumps(meta, ensure_ascii=False))
            lines.append(json.dumps(doc, ensure_ascii=False))

        payload = ("\n".join(lines) + "\n").encode("utf-8")
        st, resp = _http_ndjson(
            "POST",
            bulk_url,
            headers=_elastic_headers(cfg, ndjson=True),
            ndjson_body=payload,
            timeout_s=cfg.timeout_s,
        )
        if st not in (200, 201):
            raise RuntimeError(f"Elastic bulk failed: status={st} resp={resp}")

        items = (resp or {}).get("items", []) or []
        for it in items:
            action = it.get("index") or it.get("create") or it.get("update") or it.get("delete") or {}
            status = int(action.get("status", 0) or 0)
            if 200 <= status < 300:
                ok_total += 1
            else:
                err_total += 1
                if len(samples) < cfg.bulk_error_samples:
                    samples.append({"status": status, "_id": action.get("_id"), "error": action.get("error")})

    return ok_total, err_total, samples


# --------------------------------------------------------------------------------------
# Optional: reindex parents from Postgres -> Elasticsearch (without rereading docs)
# --------------------------------------------------------------------------------------

class ElasticReindexRequest(BaseModel):
    doc_id: Optional[str] = Field(default=None, description="Реиндекс только одного документа")
    doc_ids: Optional[List[str]] = Field(default=None, description="Список doc_id для реиндекса")
    recreate_index: bool = Field(default=False, description="Удалить и создать индекс заново")
    refresh: bool = Field(default=False, description="refresh=wait_for на bulk + _refresh в конце")
    batch_size: int = Field(default=1000, ge=50, le=5000)


class ElasticReindexResponse(BaseModel):
    doc_ids: List[str] = Field(default_factory=list)
    parents_indexed: int = 0
    bulk_errors: int = 0
    bulk_error_samples: List[Dict[str, Any]] = Field(default_factory=list)


_engine = None
_SessionLocal = None


def _get_sessionmaker():
    global _engine, _SessionLocal
    if _SessionLocal is not None:
        return _SessionLocal

    dsn = os.getenv("POSTGRES_DSN") or getattr(settings, "POSTGRES_DSN", None)
    if not dsn:
        raise RuntimeError("POSTGRES_DSN is not set (needed for Elastic reindex from DB)")

    from sqlalchemy import create_engine  # type: ignore
    from sqlalchemy.orm import sessionmaker  # type: ignore

    _engine = create_engine(dsn, pool_pre_ping=True)
    _SessionLocal = sessionmaker(bind=_engine, autoflush=False, autocommit=False)
    return _SessionLocal


def _list_doc_ids_from_db(session) -> List[str]:
    try:
        from sqlalchemy import select  # type: ignore
        rows = session.execute(select(ParentChunk.doc_id).distinct()).all()
        return [r[0] for r in rows]
    except Exception:
        rows = session.query(ParentChunk.doc_id).distinct().all()
        return [x[0] for x in rows]


def _load_parents_for_doc(session, doc_id: str) -> List[ParentChunk]:
    try:
        from sqlalchemy import select  # type: ignore
        rows = session.execute(
            select(ParentChunk).where(ParentChunk.doc_id == doc_id).order_by(ParentChunk.part_no)
        ).scalars().all()
        return list(rows)
    except Exception:
        return list(
            session.query(ParentChunk).filter(ParentChunk.doc_id == doc_id).order_by(ParentChunk.part_no).all()
        )


@router.post("/elastic/reindex_parents", response_model=ElasticReindexResponse)
def elastic_reindex_parents(req: ElasticReindexRequest) -> ElasticReindexResponse:
    t0 = time.perf_counter()
    cfg = _elastic_cfg()

    if req.recreate_index:
        cfg = ElasticConfig(
            url=cfg.url,
            index=cfg.index,
            user=cfg.user,
            password=cfg.password,
            timeout_s=cfg.timeout_s,
            refresh=req.refresh,
            recreate=True,
            bulk_actions=cfg.bulk_actions,
            bulk_error_samples=cfg.bulk_error_samples,
        )

    _elastic_ensure_index(cfg)

    SessionLocal = _get_sessionmaker()
    session = SessionLocal()
    try:
        if req.doc_id:
            doc_ids = [req.doc_id]
        elif req.doc_ids:
            doc_ids = list(dict.fromkeys(req.doc_ids))
        else:
            doc_ids = _list_doc_ids_from_db(session)

        total_indexed = 0
        total_errors = 0
        all_samples: List[Dict[str, Any]] = []

        for i, doc_id in enumerate(doc_ids, start=1):
            p0 = time.perf_counter()
            parents = _load_parents_for_doc(session, doc_id)
            if not parents:
                log.warning("[elastic_reindex] skip_empty doc_id=%s", doc_id)
                continue

            deleted = _elastic_delete_by_doc_id(cfg, doc_id)
            ok, errs, samples = _elastic_bulk_index_parents(cfg, parents, refresh_wait=req.refresh)

            total_indexed += ok
            total_errors += errs
            for s in samples:
                if len(all_samples) < cfg.bulk_error_samples:
                    all_samples.append(s)

            log.warning(
                "[elastic_reindex] doc=%d/%d doc_id=%s parents=%d deleted=%d indexed_ok=%d bulk_errors=%d (%.2fs)",
                i, len(doc_ids), doc_id, len(parents), deleted, ok, errs, time.perf_counter() - p0
            )

        if req.refresh:
            _elastic_refresh(cfg)

        log.warning(
            "[elastic_reindex] done docs=%d indexed_ok=%d bulk_errors=%d (%.2fs)",
            len(doc_ids), total_indexed, total_errors, time.perf_counter() - t0
        )

        return ElasticReindexResponse(
            doc_ids=doc_ids,
            parents_indexed=total_indexed,
            bulk_errors=total_errors,
            bulk_error_samples=all_samples,
        )
    finally:
        try:
            session.close()
        except Exception:
            pass


# --------------------------------------------------------------------------------------
# Main ingest endpoint
# --------------------------------------------------------------------------------------

@router.post("/ingest", response_model=IngestResponse)
def ingest(req: IngestRequest) -> IngestResponse:
    t0 = time.perf_counter()
    repo = Repository()

    parent_chars = int(getattr(settings, "PARENT_CHARS", 2500))

    enable_qdrant_indexing = _truthy(os.getenv("ENABLE_INDEXING", "0"))
    enable_elastic_indexing = _truthy(os.getenv("ENABLE_ELASTIC_INDEXING", os.getenv("ENABLE_INDEXING", "0")))

    folder = Path(req.folder_path)
    if not folder.is_absolute():
        folder = Path("/app") / folder

    if not folder.exists() or not folder.is_dir():
        raise HTTPException(status_code=400, detail=f"folder_path not found: {folder}")

    exts = {e.lower() for e in req.extensions}
    paths = list(folder.rglob("*")) if req.recursive else list(folder.glob("*"))
    files = [
        p for p in paths
        if p.is_file()
        and p.suffix.lower() in exts
        and not p.name.startswith("~$")
    ]
    files.sort(key=lambda p: str(p).lower())

    child_cfg = ChildChunkConfig(
        max_len=int(os.getenv("CHILD_MAX_LEN", "600")),
        overlap=int(os.getenv("CHILD_OVERLAP", "80")),
        min_len=int(os.getenv("CHILD_MIN_LEN", "120")),
    )

    qdrant_batch = max(32, min(int(os.getenv("QDRANT_BATCH", "512")), 4096))
    log.warning("[ingest] qdrant_batch=%d", qdrant_batch)

    elastic_cfg: ElasticConfig | None = None
    if enable_elastic_indexing:
        elastic_cfg = _elastic_cfg()
        _elastic_ensure_index(elastic_cfg)
        log.warning(
            "[ingest] elastic_ready url=%s index=%s refresh=%s recreate=%s bulk_actions=%d",
            elastic_cfg.url, elastic_cfg.index, elastic_cfg.refresh, elastic_cfg.recreate, elastic_cfg.bulk_actions
        )

    qdrant_store: QdrantStore | None = get_qdrant_store() if enable_qdrant_indexing else None

    log.warning(
        "[ingest] start folder=%s files=%d parent_chars=%d qdrant=%s elastic=%s",
        str(folder), len(files), parent_chars, enable_qdrant_indexing, enable_elastic_indexing
    )

    indexed_files = 0
    parents_total = 0
    children_db_total = 0
    skipped_files = 0
    errors: List[str] = []

    for fp in files:
        file_t0 = time.perf_counter()
        try:
            if fp.suffix.lower() != ".docx":
                skipped_files += 1
                continue

            log.warning("[ingest] file=%s", fp.name)

            text = _extract_text_from_docx(fp)
            if not text:
                skipped_files += 1
                log.warning("[ingest] skip empty_text file=%s", fp.name)
                continue

            blob = fp.read_bytes()
            doc_id = _stable_id("doc", blob)
            title, doc_number, doc_date = _guess_doc_meta(text, fallback_title=fp.name)

            repo.delete_document(doc_id)
            repo.upsert_document(doc_id=doc_id, title=title, doc_number=doc_number, doc_date=doc_date)

            parents: List[ParentChunk] = []
            for part_no, chunk_text in _parent_chunk(text, parent_chars):
                parent_payload = f"{doc_id}:{part_no}:{chunk_text[:200]}".encode("utf-8", errors="ignore")
                parent_id = _stable_id("parent", parent_payload)
                parents.append(
                    ParentChunk(
                        parent_id=parent_id,
                        doc_id=doc_id,
                        section_path="",
                        anchor="",
                        title=title,
                        part_no=part_no,
                        content=chunk_text,
                    )
                )

            pw = repo.upsert_parents(parents)
            parents_total += pw
            log.warning("[ingest] parents_written=%d parents_in_mem=%d doc_id=%s", pw, len(parents), doc_id)

            if enable_elastic_indexing and parents:
                assert elastic_cfg is not None
                t_es0 = time.perf_counter()
                try:
                    deleted = _elastic_delete_by_doc_id(elastic_cfg, doc_id)
                    ok, bulk_errs, samples = _elastic_bulk_index_parents(
                        elastic_cfg, parents, refresh_wait=elastic_cfg.refresh
                    )
                    if bulk_errs:
                        log.warning("[ingest] elastic_bulk_errors count=%d samples=%s", bulk_errs, samples)
                        errors.append(f"{fp.name}: Elastic bulk errors: {bulk_errs} samples={samples}")
                    log.warning(
                        "[ingest] elastic_ok deleted=%d indexed_ok=%d bulk_errors=%d (%.2fs)",
                        deleted, ok, bulk_errs, time.perf_counter() - t_es0
                    )
                except Exception as e:
                    log.warning("[ingest] elastic_failed file=%s err=%r", fp.name, e)
                    errors.append(f"{fp.name}: ElasticError: {e}")

            t_children0 = time.perf_counter()
            log.warning("[ingest] children_build_start parents=%d", len(parents))
            children = _build_children(parents, child_cfg)
            log.warning(
                "[ingest] children_build_done children=%d (%.2fs)",
                len(children), time.perf_counter() - t_children0
            )

            t_db0 = time.perf_counter()
            cw = repo.upsert_children(children)
            children_db_total += cw
            log.warning("[ingest] children_db_written=%d (%.2fs)", cw, time.perf_counter() - t_db0)

            if enable_qdrant_indexing:
                assert qdrant_store is not None
                try:
                    t_del0 = time.perf_counter()
                    deleted = qdrant_store.delete_by_doc_id(doc_id, wait=True)
                    log.warning(
                        "[ingest] qdrant_delete_by_doc_id deleted=%s (%.2fs) doc_id=%s",
                        deleted, time.perf_counter() - t_del0, doc_id
                    )

                    if children:
                        t_up0 = time.perf_counter()
                        qw = qdrant_store.upsert_children(children, batch_size=qdrant_batch)
                        log.warning(
                            "[ingest] qdrant_ok written=%d (%.2fs) collection=%s doc_id=%s",
                            qw, time.perf_counter() - t_up0, qdrant_store.collection, doc_id
                        )
                    else:
                        log.warning("[ingest] qdrant_skip no_children doc_id=%s", doc_id)

                except Exception as e:
                    log.warning("[ingest] qdrant_failed file=%s doc_id=%s err=%r", fp.name, doc_id, e)
                    errors.append(f"{fp.name}: QdrantError: {e}")

            indexed_files += 1
            log.warning("[ingest] file_done=%s (%.2fs)", fp.name, time.perf_counter() - file_t0)

        except Exception as e:
            log.warning("[ingest] file_failed=%s err=%r", fp.name, e)
            errors.append(f"{fp.name}: {type(e).__name__}: {e}")

    if enable_elastic_indexing and elastic_cfg is not None and elastic_cfg.refresh:
        _elastic_refresh(elastic_cfg)

    log.warning(
        "[ingest] done indexed=%d parents=%d children_db=%d skipped=%d (%.2fs)",
        indexed_files, parents_total, children_db_total, skipped_files, time.perf_counter() - t0
    )

    return IngestResponse(
        indexed_files=indexed_files,
        parents=parents_total,
        children=children_db_total,
        skipped_files=skipped_files,
        errors=errors,
    )
